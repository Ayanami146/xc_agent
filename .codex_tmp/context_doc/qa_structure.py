from __future__ import annotations

import hashlib
import zipfile
from pathlib import Path

from docx import Document


DOCX = Path(
    r"C:\work_learn\XinChuang_pc\.codex_tmp\context_doc\基于LangChain_LangGraph官方组件的智能体上下文开发指南.docx"
)

document = Document(DOCX)
all_text = "\n".join(
    [paragraph.text for paragraph in document.paragraphs]
    + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
)

required = [
    "AsyncShallowRedisSaver",
    "SummarizationMiddleware",
    "MongoDBStore",
    "xc:agent:context:persist:v1",
    'durability="sync"',
    "messages_to_dict",
    "messages_from_dict",
    "【系统说明：用户已取消本轮回答，以上内容可能不完整或不正确。】",
    "langgraph-checkpoint-redis>=0.5.2,<0.6",
    "langgraph-store-mongodb>=0.3,<0.4",
    "https://docs.langchain.com/oss/python/langgraph/persistence",
    "https://github.com/redis-developer/langgraph-redis",
]
missing = [item for item in required if item not in all_text]
assert not missing, f"缺少关键内容：{missing}"
assert 'content = f"{partial}\\n\\n{CANCELLED_NOTICE}"' in all_text

with zipfile.ZipFile(DOCX) as archive:
    assert archive.testzip() is None, "DOCX ZIP 结构损坏"

print(
    f"paragraphs={len(document.paragraphs)} tables={len(document.tables)} "
    f"chars={len(all_text)} size={DOCX.stat().st_size}"
)
print(f"sha256={hashlib.sha256(DOCX.read_bytes()).hexdigest().upper()}")
print("required_phrases=OK exact_cancel_notice=OK escaped_newline=OK zip=OK")
