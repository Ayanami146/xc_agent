from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document


def extract(path: Path) -> dict:
    doc = Document(path)
    paragraphs = []
    for index, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            paragraphs.append(
                {
                    "index": index,
                    "style": paragraph.style.name if paragraph.style else None,
                    "text": text,
                }
            )

    tables = []
    for table_index, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append({"index": table_index, "rows": rows})

    return {
        "path": str(path),
        "paragraphs": paragraphs,
        "tables": tables,
        "sections": len(doc.sections),
    }


if __name__ == "__main__":
    payload = [extract(Path(argument)) for argument in sys.argv[1:]]
    print(json.dumps(payload, ensure_ascii=False, indent=2))
