"""生成“取消运行与上下文一致性”学习问题记录。"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(
    r"C:\work_learn\XinChuang_pc\.codex_tmp\context_issue_note"
    r"\问题记录_取消运行与上下文一致性.docx"
)

NAVY = "17365D"
BLUE = "2F75B5"
TEXT = "202124"
MUTED = "667085"
PALE_BLUE = "EAF2F8"
PALE_GREEN = "E8F5EE"
PALE_AMBER = "FFF4DD"
PALE_GRAY = "F3F5F7"
BORDER = "C9D3DF"


def set_cell_shading(cell, fill: str) -> None:
    """设置单元格底色。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    """固定 Word 单元格宽度，避免不同渲染器自动挤压。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    """设置表格总宽度、列网格和每个单元格宽度。"""
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])


def set_run_font(run, name: str = "Microsoft YaHei", size: float = 10.5) -> None:
    """同时设置中西文字体。"""
    run.font.name = name
    run.font.size = Pt(size)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), name)


def style_paragraph(paragraph, *, after: float = 6, line: float = 1.35) -> None:
    """应用正文段落节奏。"""
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_heading(document: Document, text: str, level: int = 1) -> None:
    """添加统一蓝色标题。"""
    paragraph = document.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size=17 if level == 1 else 12.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE if level == 1 else NAVY)


def add_body(document: Document, text: str, *, bold_prefix: str | None = None) -> None:
    """添加正文；可选择加粗开头标签。"""
    paragraph = document.add_paragraph()
    style_paragraph(paragraph)
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first)
        first.font.bold = True
        rest = paragraph.add_run(text[len(bold_prefix) :])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)


def add_bullet(document: Document, text: str) -> None:
    """使用 Word 内置列表样式添加项目符号。"""
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.18)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run_font(run)


def create_numbering(document: Document) -> int:
    """创建一组从 1 开始的真实 Word 编号，并返回 numId。"""
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(item.get(qn("w:abstractNumId")))
        for item in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(item.get(qn("w:numId")))
        for item in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    level_jc = OxmlElement("w:lvlJc")
    level_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "432")
    ind.set(qn("w:hanging"), "288")
    p_pr.append(ind)
    level.extend([start, num_fmt, level_text, level_jc, p_pr])
    abstract.append(level)

    # OOXML 要求所有 abstractNum 位于 num 实例之前。
    # 如果直接追加到文件末尾，Word 可能把其他列表样式错误解释为连续编号。
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(numbering.index(first_num), abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_number(document: Document, text: str, *, num_id: int) -> None:
    """使用指定的真实 Word 编号组添加步骤。"""
    paragraph = document.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_element])
    p_pr.append(num_pr)
    paragraph.paragraph_format.left_indent = Inches(0.3)
    paragraph.paragraph_format.first_line_indent = Inches(-0.2)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    set_run_font(run)


def add_callout(document: Document, title: str, body: str, fill: str) -> None:
    """添加提示框。"""
    table = document.add_table(rows=1, cols=1)
    # 提示框的标题和正文必须作为一个整体排版，避免标题落在上一页底部、
    # 正文被拆到下一页，影响阅读连续性。
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(4)
    title_run = paragraph.add_run(title)
    set_run_font(title_run, size=11)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor.from_string(NAVY)
    body_paragraph = cell.add_paragraph()
    body_paragraph.paragraph_format.space_after = Pt(2)
    body_paragraph.paragraph_format.line_spacing = 1.35
    body_run = body_paragraph.add_run(body)
    set_run_font(body_run)


def add_code(document: Document, code: str) -> None:
    """添加紧凑伪代码块。"""
    table = document.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_GRAY)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(code.strip())
    set_run_font(run, name="Consolas", size=8.5)


def add_page_number(paragraph) -> None:
    """插入动态页码字段。"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    suffix = paragraph.add_run(" 页")
    set_run_font(suffix, size=9)


def build() -> None:
    """创建问题记录文档。"""
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    header = section.header.paragraphs[0]
    header.text = "上下文系统 · 学习问题记录"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        set_run_font(run, size=8.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(MUTED)

    add_page_number(section.footer.paragraphs[0])

    eyebrow = document.add_paragraph()
    eyebrow.paragraph_format.space_after = Pt(6)
    run = eyebrow.add_run("暂缓实现 / 后续专题")
    set_run_font(run, size=10)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(7)
    title_run = title.add_run("取消运行与上下文一致性问题")
    set_run_font(title_run, size=24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor.from_string(NAVY)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle.add_run(
        "当用户取消智能体后立即发送新请求，如何保证上一轮取消上下文已经安全提交？"
    )
    set_run_font(subtitle_run, size=12)
    subtitle_run.font.color.rgb = RGBColor.from_string(BLUE)

    meta = document.add_table(rows=4, cols=2)
    set_table_geometry(meta, [2200, 7160])
    values = [
        ("记录日期", "2026-08-28"),
        ("当前状态", "已记录，暂缓实现"),
        ("讨论阶段", "完成上下文系统第一版后重新评审"),
        ("关联模块", "取消注册表、SessionRunRegistry、LangGraph、ContextRepository、SSE"),
    ]
    for row_index, (label, value) in enumerate(values):
        left, right = meta.rows[row_index].cells
        set_cell_shading(left, PALE_BLUE)
        for cell, text_value, bold in ((left, label, True), (right, value, False)):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(text_value)
            set_run_font(run, size=9.5)
            run.font.bold = bold

    document.add_paragraph()
    add_callout(
        document,
        "当前学习决策",
        "先完成原定的上下文系统第一版，不在中途引入取消收尾状态机、取消轮次持久化和恢复机制。该问题单独保留，等第一版闭环跑通后再实现，避免代码结构与学习顺序同时变化。",
        PALE_GREEN,
    )

    add_heading(document, "一、问题是怎样出现的")
    add_body(
        document,
        "同一个会话中，用户正在接收智能体的流式回答。用户点击取消后，模型虽然停止继续生成，但本轮可能已经完成了上下文加载、RAG 检索或工具调用，也可能已经产生部分回答。此时如果用户立即发送下一条消息，新请求可能在上一轮取消信息尚未保存时读取旧快照。",
    )
    add_code(
        document,
        """
会话当前快照：revision = 3

请求 A：正在运行
  -> 用户点击取消
  -> 模型停止输出
  -> 取消上下文尚未写入 MongoDB

请求 B：立即开始
  -> 仍然读取 revision = 3
  -> 不知道请求 A 已被取消，也读不到 A 已确认的工具/RAG 信息
        """,
    )
    add_body(
        document,
        "潜在结果：请求 B 基于过期上下文生成；请求 A 和 B 在保存时出现 revision 冲突；或者为了避免冲突而丢弃请求 A 的取消上下文。",
        bold_prefix="潜在结果：",
    )

    add_heading(document, "二、这个问题真正难在哪里")
    add_bullet(document, "“取消请求已经受理”不等于“智能体运行已经完成”。")
    add_bullet(document, "取消时可能有用户问题、工具结果和 RAG 证据，但模型回答通常只有半截。")
    add_bullet(document, "如果取消轮次需要参与后续模型上下文，就必须明确它是否增加 revision。")
    add_bullet(document, "如果 MongoDB 保存较慢或失败，必须决定新请求是等待、报错还是进入恢复状态。")
    add_bullet(document, "前端禁用发送只能改善体验，后端仍需防止重复请求、重试和多页面并发。")

    add_callout(
        document,
        "关键认识",
        "revision 可以防止旧快照覆盖新快照，但不能单独保证第二个请求在生成前已经看到第一个请求的取消上下文。要解决时序问题，还需要会话级串行控制或原子化的运行占用机制。",
        PALE_AMBER,
    )

    add_heading(document, "三、当前第一版保持什么边界")
    first_version_num = create_numbering(document)
    add_number(document, "先完成 Redis 优先、MongoDB 回源、正常轮次保存、上下文压缩和 revision 乐观锁。", num_id=first_version_num)
    add_number(document, "取消功能继续使用现有最小流程，不在这一阶段扩展取消轮次的数据结构。", num_id=first_version_num)
    add_number(document, "不把流式生成的半截 assistant_text 当作一条完整模型回答保存。", num_id=first_version_num)
    add_number(document, "不在第一版同时引入消息队列、分布式锁、补偿任务或 RECOVERY_REQUIRED 状态。", num_id=first_version_num)
    add_number(document, "第一版通过测试并完成实际 Redis/MongoDB 验证后，再回到本文讨论取消一致性。", num_id=first_version_num)

    add_heading(document, "四、后续重新讨论时的候选方向")
    add_body(
        document,
        "以下内容只是后续设计线索，不是当前版本的实施要求。届时应先重新检查实际代码，再决定是否采用。",
    )
    future_num = create_numbering(document)
    add_number(document, "取消接口只返回 CANCEL_REQUESTED，不立即释放会话运行权。", num_id=future_num)
    add_number(document, "模型停止后进入 FINALIZING_CANCEL，整理可保留的用户问题、已完成工具结果和 RAG 信息。", num_id=future_num)
    add_number(document, "取消轮次使用 status=cancelled；partial assistant output 不作为完整回答进入上下文。", num_id=future_num)
    add_number(document, "取消快照提交 MongoDB 成功后再更新 Redis，并在需要时增加 revision。", num_id=future_num)
    add_number(document, "只有取消收尾完成后才释放 SessionRunRegistry，并向前端发送最终 RUN_CANCELLED。", num_id=future_num)
    add_number(document, "如果持久化失败，评估有限重试、会话阻塞、补偿任务或恢复状态。", num_id=future_num)

    add_code(
        document,
        """
未来可能的状态机（非当前实现）

RUNNING
  -> CANCEL_REQUESTED
  -> FINALIZING_CANCEL
  -> persist_cancelled_context
  -> CANCELLED
  -> release_session

新请求只有在 release_session 之后才能进入 load_context。
        """,
    )

    add_heading(document, "五、后续必须回答的问题")
    questions = [
        "取消轮次只作为运行审计保存，还是也要参与后续模型上下文？",
        "用户问题是否保留？已经成功的工具结果和 RAG 结果保留到什么粒度？",
        "取消轮次是否增加 context revision？什么时刻视为提交成功？",
        "取消时已经输出给用户的半截回答是否保存到 MySQL 展示记录？模型上下文是否完全忽略它？",
        "MongoDB 暂时不可用时，会话应当阻塞多久？是否需要补偿任务？",
        "客户端断开、显式取消、模型异常三种终止路径是否采用同一保存策略？",
        "单进程 SessionRunRegistry 何时升级为 Redis 分布式锁或后端队列？",
    ]
    for question in questions:
        add_bullet(document, question)

    add_heading(document, "六、未来实现后的验收场景")
    checks = [
        "取消请求返回后，前端在取消收尾完成前不能发送下一条消息。",
        "正常取消不会把半截模型回答保存为 completed 上下文。",
        "已经完成的工具/RAG 信息能按最终策略保留或明确丢弃。",
        "取消上下文提交成功后，新请求读取到最新 revision。",
        "重复取消、重复 requestId 和客户端断开不会重复增加 revision。",
        "MongoDB 保存失败时不会静默释放会话并让新请求读取旧快照。",
        "多页面或并发请求仍然只能有一个请求占用同一会话。",
    ]
    for check in checks:
        add_bullet(document, check)

    add_callout(
        document,
        "何时回来处理",
        "当上下文系统第一版完成以下闭环后再继续：正常两轮对话可从 Redis/MongoDB 正确加载；revision 按轮增加；Redis 故障可回源 MongoDB；上下文压缩与现有测试通过。",
        PALE_BLUE,
    )

    document.core_properties.title = "取消运行与上下文一致性问题记录"
    document.core_properties.subject = "上下文系统后续问题"
    document.core_properties.keywords = "取消运行, 上下文, revision, 并发, 暂缓实现"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
